import whisper
import pymysql
import os
import warnings
warnings.filterwarnings("ignore", category=UserWarning)
from datetime import datetime
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from openpyxl import Workbook

# 1. Whisper 모델 로드
model = whisper.load_model("small")

# 2. 오디오 파일 경로
audio_path = r"C:\Users\user\voice2text_project\voice\001_sd.m4a"
audio_filename = os.path.basename(audio_path)

# 3. 음성 텍스트 추출
result = model.transcribe(audio_path, language="ko")
full_text = result["text"]

# 4. 텍스트 요약 (간단 요약 - 첫 50자 기준)
summary_text = full_text[:50] + "..." if len(full_text) > 50 else full_text

# 5. 파일 형식 선택 (예: "docx", "pdf", "xlsx", "txt")
file_format = "xlsx"  # 원하는 형식으로 변경 (예: "docx", "pdf", "xlsx", "txt")

# 6. 타임스탬프 생성
now = datetime.now()
timestamp = now.strftime("%Y%m%d_%H%M%S")

# 7. 파일 경로 설정
save_directory = r"C:\Users\user\voice2text_project\doc"
os.makedirs(save_directory, exist_ok=True)

# 8. 저장할 파일 경로 설정
filename = f"voice_{timestamp}.{file_format}"
file_path = os.path.join(save_directory, filename)

# 9. DOCX 저장
def save_as_docx():
    doc = Document()
    doc.add_heading("요약", level=1)
    doc.add_paragraph(summary_text)
    doc.add_heading("전체 텍스트", level=1)
    doc.add_paragraph(full_text)
    doc.save(file_path)

# 10. PDF 저장
def save_as_pdf():
    c = canvas.Canvas(file_path, pagesize=letter)
    c.drawString(100, 750, "요약")
    c.drawString(100, 730, summary_text)
    c.drawString(100, 710, "전체 텍스트")
    c.drawString(100, 690, full_text)
    c.save()

# 11. XLSX 저장
def save_as_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Voice Summary"
    ws['A1'] = "요약"
    ws['A2'] = summary_text
    ws['A3'] = "전체 텍스트"
    ws['A4'] = full_text
    wb.save(file_path)

# 12. TXT 저장
def save_as_txt():
    with open(file_path, 'w', encoding='utf-8') as file:
        file.write("요약\n")
        file.write(summary_text + "\n\n")
        file.write("전체 텍스트\n")
        file.write(full_text)

# 13. 파일 형식에 맞는 저장 함수 호출
if file_format == "docx":
    save_as_docx()
elif file_format == "pdf":
    save_as_pdf()
elif file_format == "xlsx":
    save_as_xlsx()
elif file_format == "txt":
    save_as_txt()

print(f"파일이 저장되었습니다: {file_path}")

# 14. DB 연결
conn = pymysql.connect(
    host='49.50.160.162',
    user='dbcvt',
    password='dbcvt1234',
    database='dbcvt',
    port=3306,
    charset='utf8mb4',
    autocommit=True
)
cursor = conn.cursor()

# 15. DB INSERT
cursor.execute("""
    INSERT INTO tb_voice_text (
        created_at,
        summary_text,
        full_text,
        audio_filename,
        document_filename,
        format
    ) VALUES (NOW(), %s, %s, %s, %s, %s)
""", (
    summary_text,
    full_text,
    audio_filename,
    filename,
    file_format
))

print("✅ tb_voice_text 테이블에 정상 저장되었습니다.")

# 16. 연결 종료
cursor.close()
conn.close()
