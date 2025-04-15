import os
import datetime
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
import whisper
from docx import Document
from docx.shared import Pt
import pymysql
import openai

app = Flask(__name__)
CORS(app)

# 디렉토리 설정
UPLOAD_FOLDER = 'uploads'
DOC_FOLDER = 'doc'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(DOC_FOLDER, exist_ok=True)

# Whisper 모델 로드 (기본 base 사용, 필요시 변경 가능)
model = whisper.load_model("base")

# OpenAI API 키 설정 (요약용)
openai.api_key = "YOUR_OPENAI_API_KEY"  # <- 실제 키 입력 필요

# MariaDB 설정
db = pymysql.connect(
    host='localhost',
    user='your_user',
    password='your_password',
    database='your_db',
    charset='utf8mb4'
)

@app.route('/transcribe', methods=['POST'])
def transcribe():
    file = request.files['audio']
    if not file:
        return jsonify({'error': 'No file uploaded'}), 400

    # 파일 저장
    timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    audio_filename = f"voice_{timestamp}.m4a"
    audio_path = os.path.join(UPLOAD_FOLDER, audio_filename)
    file.save(audio_path)

    # Whisper로 텍스트 변환
    result = model.transcribe(audio_path, language='ko')
    full_text = result["text"]

    # 요약 생성 (OpenAI GPT-3.5)
    try:
        summary_prompt = f"다음 내용을 한국어로 짧게 요약해줘:\n\n{full_text}"
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": summary_prompt}]
        )
        summary_text = response['choices'][0]['message']['content'].strip()
    except Exception as e:
        summary_text = "요약 생성 실패"

    # Word 문서 생성
    document = Document()
    document.add_heading('요약', level=1)
    summary_paragraph = document.add_paragraph(summary_text)
    summary_paragraph.style.font.size = Pt(12)

    document.add_heading('전체 텍스트', level=1)
    full_paragraph = document.add_paragraph(full_text)
    full_paragraph.style.font.size = Pt(11)

    doc_filename = f"voice_summary_{timestamp}.docx"
    doc_path = os.path.join(DOC_FOLDER, doc_filename)
    document.save(doc_path)

    # DB에 저장
    try:
        with db.cursor() as cursor:
            sql = """
                INSERT INTO tb_voice_text (
                    created_at,
                    summary_text,
                    full_text,
                    audio_filename,
                    document_filename,
                    format
                ) VALUES (NOW(), %s, %s, %s, %s, %s)
            """
            cursor.execute(sql, (
                summary_text,
                full_text,
                audio_filename,
                doc_filename,
                'docx'
            ))
            db.commit()
    except Exception as e:
        print("DB Insert Error:", e)

    return jsonify({
        'summary': summary_text,
        'text': full_text,
        'audio_filename': audio_filename,
        'document_filename': doc_filename
    })


@app.route('/doc/<filename>')
def download_doc(filename):
    return send_from_directory(DOC_FOLDER, filename, as_attachment=True)


if __name__ == '__main__':
    app.run(port=5000, debug=True)
